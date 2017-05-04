# -*- coding: utf-8 -*-
"""
Created on Mon Jul 18 12:24:13 2016

@author: babou
"""
import pandas as pd
import glob
import nltk
import random

from nltk.corpus import stopwords
from nltk.stem import SnowballStemmer
from docx import Document
from sklearn.preprocessing import LabelEncoder
from textblob import TextBlob
from textblob_fr import PatternTagger, PatternAnalyzer


#target_list = ['X\xe2\x80\xa6', 'Y\xe2\x80\xa6', 'Z\xe2\x80\xa6', 'X']
stemmer = SnowballStemmer("french")
stopword_fr = [word for word in stopwords.words('french')]

target_dict = {u'X\xe2\x80\xa6' : 'X', u'X' : 'X', u'X..' : 'X', u'X.' : 'X',
               u'X\u2026' : 'X', u'M.X' : 'X', u'M.X\u2026' : 'X',
               u'Y\xe2\x80\xa6' : 'Y', u'Y' : 'Y', u'Y..' : 'Y', u'Y.' : 'Y',
               u'Y\u2026' : 'Y', u'M.Y' : 'Y', u'M.Y\u2026' : 'Y',
               u'Z\xe2\x80\xa6': 'Z', u'Z' : 'Z', u'Z..' : 'Z', u'Z.' : 'Z',
               u'Z\u2026' : 'Z', u'M.Z' : 'Z', u'M.Z\u2026' : 'Z'}

mister_list = [u'M', u'M.', u'Madame', u'Mme', u'Monsieur', u'Dr', u'Monsieur', u'MM']

#punctuations = '''!()-[]{};:'"\,<>./?@#$%^&*_~'''

#################################################
###                 FUNCTIONS                 ###
#################################################


#def get_target_name(sentence):
#    for i in range(len(sentence.split(' '))):
#        w = sentence.split(' ')[i].replace(',', '').replace('.', '')
#        if w in target_list:
#            return target_dict.keys.index(w)


def get_target(sentence):
    for i in range(len(sentence.split(' '))):
        w = sentence.split(' ')[i].replace(',', '').replace('.', '')
        if w in target_dict.keys():
            return 1

#def remove_ponctuation(sentence):
#    sentence_with_no_punct = ""
#    for char in sentence:
#       if char not in punctuations:
#           sentence_with_no_punct = sentence_with_no_punct + char
#    return sentence_with_no_punct



#################################################
###                 GENERAL                   ###
#################################################

path_files = glob.glob('data/jugements CRC LÇgifrance/*')
doc = Document(path_files[0])

documents_df = pd.DataFrame()
document_temp = pd.DataFrame()


for path_file in path_files:
    print "Loading file : " + path_file.split('/')[-1]
    if path_file.split('.')[-1] == "docx":
        doc = Document(path_file)
        paragraphs = doc.paragraphs
        my_document = []
        for paragraph in paragraphs:
            my_document.append({'doc_name' : path_file.split('/')[-1],
                                'paragraph' : paragraph.text})

        document_temp = pd.DataFrame(my_document)
        documents_df = pd.concat([documents_df, document_temp])
    else:
        print "Error - Not a docx file : " + path_file.split('/')[-1]

#documents_df['paragraph_no_punct'] = documents_df['paragraph'].apply(lambda x: remove_ponctuation(x))
documents_df.reset_index(inplace=True)
documents_df['is_target'] = documents_df['paragraph'].apply(lambda x: get_target(x))
documents_df['is_target'].fillna(0, inplace=True)


print documents_df.is_target.value_counts()
print documents_df.groupby('doc_name')['is_target'].sum()


# Old
#    tags = tagger.TagText(row.['paragraph'])
#    tags2 = treetaggerwrapper.make_tags(tags)
#    for w in tags2:
#        word_list.append({'word' : w[0],
#                          'type' : w[1],
#                            'lemma' : w[2]})

print "-"*54
print " PREPROCESSING..."

word_list = []
context_word =[]
for idx , row in documents_df.T.iteritems():
    tokenized = nltk.word_tokenize(row['paragraph'])
    word_list.extend(tokenized)
    # Some context paragraph / doc
    for len_token in tokenized:
        context_word.append({'doc_name' : row['doc_name'],
                             'paragraph_nb' : row['index']})

# My bag of word
word_df = pd.DataFrame(word_list, columns=['word'])

# My context
context_df = pd.DataFrame(context_word)

# Merging context & bad of word
word_df = pd.concat([word_df, context_df], axis=1)


## Loading annexe files (Firstaname / Names etc;..)
# Reading French's firstnames file
firstname_df = pd.read_csv('data/prenom_clean.csv', encoding='utf-8',
                           header=None, names = ["firstname"])
# Use Majuscule in first carac
firstname_df['firstname'] = firstname_df['firstname'].apply(lambda x: x.title())
firstname_list = firstname_df.firstname.tolist()


# Reading foreign's firstnames file
foreign_firstname_df =  pd.read_csv('data/foreign_fistname_clean.csv', encoding='utf-8',
                             header=None, names = ["firstname"])
foreign_firstname_list = foreign_firstname_df.firstname.tolist()

# Name list
name_df = pd.read_csv('data/top_8k_name.csv', encoding='utf-8')
name_list = name_df.name.tolist()


#################################################
###         Simulate Name & Firstame          ###


word_df['is_target'] = word_df['word'].apply(lambda x: 1 if x in target_dict.keys() else 0)
# To make some difference beetwen a name is_target and other target. To delete btw
word_df['is_name'] = 0
word_df.loc[word_df['is_target'] == 1, 'is_name'] = 1


## Insert random first name in place of 'X', Mr 'X...' ...
#

# /!\ Delete when true data
def get_random_firstname():
    """
    Use a first random (0/1) to chose between French's firstname or foreign
    And a random firstname in the random chosen dataset
    """
    # first random to chose the dataset
    first_random = random.randint(0,1)
    # French
    if first_random == 0:
        random_idx = random.randint(0, 2046) # len de firstname_list
        return firstname_list[random_idx]
    # foreign
    else:
        random_idx = random.randint(0, 36114) # len de firstname_list
        return foreign_firstname_list[random_idx]

def get_random_name():
    """
    Return a random name (CAPS or tiltle)
    """
    first_random = random.randint(0,1)
    random_idx = random.randint(0, 7999) # len de name_list
    if first_random == 0:
        return name_list[random_idx]
    else:
        return name_list[random_idx].title() 

def insert_row(x, idx):
    """
    Return a Dataframe with :
    - Random Firstname (French or Foreigner)
    - is_target = 1
    -
    """
    line = {'word' : get_random_firstname() ,
            'doc_name' :  x['doc_name'],
            'paragraph_nb' :  x['paragraph_nb'],
            'is_target' :  1,
            'is_firstname' :  1,
            'word_shift_1b' :  '',
            'is_firstname_1b' :  '',
            'is_firstname_1b_shift' :  '',
            'add_row' : 1,#
            'is_target_1b' : ''}
    return pd.DataFrame(line, index=[idx])


# Replace X.. Y by Radnom Name
word_df.loc[word_df['is_target'] ==1, 'word'] = word_df['word'].apply(lambda x: get_random_name())
word_df['admin_name'] = 0 # Admin stuff 
word_df.loc[word_df['is_target'] ==1, 'admin_name'] = 1 # Admin stuff 


# Check if word is Firstname then is_target :
word_df['is_firstname'] = word_df['word'].apply(lambda x: 1 if x.title() in firstname_list else 0)
word_df['len_word'] = word_df['word'].apply(lambda x: len(x))
#  We need to constrait word with forgein Firstanme list (to much false positive)
word_df.loc[((word_df.len_word) > 3) & (word_df['word'].isin(foreign_firstname_list)), 'is_firstname' ] = 1
word_df = word_df.drop('len_word', axis=1)

word_df['word_shift_1b'] = word_df.groupby(['doc_name'])['word'].apply(lambda x: x.shift(1))
word_df['is_firstname_1b'] = 0
word_df.loc[(word_df['is_target'] == 1) & (word_df['word_shift_1b'].str.title().isin(firstname_list)), 'is_firstname_1b'] = 1
word_df['is_firstname_1b_shift'] = word_df['is_firstname_1b'].shift(-1)
word_df.loc[(word_df['is_firstname_1b_shift'] == 1), 'is_target'] = 1



# Random Insert Random Firstname before Name if no firstname detect
print "  INSERT ROWS FOR SIMULATE FIRSTNAME"

word_df['add_row'] = 0
word_df['is_target_1b'] = word_df['is_target'].shift(-1)

word_df['is_firstname_1a'] = word_df['is_firstname'].shift(-1)

i = 0
for idx, row in word_df[(word_df.is_target_1b == 1) & (word_df.is_firstname == 0)
                        & (word_df.is_firstname_1a == 0)].iterrows():
    # Add some random (sometime Firstnane sometime no...)
    if random.randint(0, 1) == 1:
        print str(idx+i)
        # Si la row suivante est target et que ce n'est pas un Firstname
        # et que la row actuelle n'est pas un Firstname ==> Add Random Firstname
        line = insert_row(row, idx)
        word_df = pd.concat([word_df.ix[:idx + i ], line, word_df.ix[idx+1 + i:]]).reset_index(drop=True)
        i+=1
    

#Delete old features
word_df = word_df.drop(['is_firstname_1b', 'is_firstname_1b_shift', 'is_target_1b', 
                        'word_shift_1b', 'is_name', 'is_firstname_1a'], axis=1)

word_df = word_df[word_df.word != u'\u2026'] # To clean anonymasation from dataset and bad ML
word_df = word_df[word_df.word != u"..."]
word_df = word_df[word_df.word != u".."]
word_df = word_df.reset_index(drop=True)


##################################################
####             Features Engi                 ###

# to have granularite
word_df['temp_count'] = 1

# Cumulative sum of word by paragraph
word_df['paragraph_cum_word' ] = word_df.groupby(['doc_name', 'paragraph_nb'])['temp_count'].cumsum()

# Cumulative sum of word by senstence end by ";" or "."
word_df['temp_count'] = 1
## Create a bool a each end of sentence
word_df["end_point"] = word_df.word.apply(lambda x: 1 if x in [";", "."] else 0)

word_df['end_point_cum' ] = word_df.groupby(['doc_name'])['end_point'].cumsum()
word_df['end_point_cum_word' ] = word_df.groupby(['doc_name', 'end_point_cum'])['temp_count'].cumsum()
#word_df = word_df.drop(['temp_count', 'end_point', 'end_point_cum'], axis=1)

# Cumulative sum of word by senstence end by ","
## Create a bool a each end of sentence
word_df["end_comma"] = word_df.word.apply(lambda x: 1 if x in [","] else 0)
## If end of sentence "." & ";" then end of comma to
word_df.loc[word_df['end_point'] == 1, "end_comma"] = 1

word_df['end_comma_cum' ] = word_df.groupby(['doc_name'])['end_comma'].cumsum()
word_df['end_comma_cum_word' ] = word_df.groupby(['doc_name', 'end_comma_cum'])['temp_count'].cumsum()

# Del temp preprocessing features
word_df = word_df.drop(['temp_count', 'end_comma', 'end_comma_cum',
                       'end_point', 'end_point_cum'], axis=1)



word_df['is_stopword'] = word_df['word'].apply(lambda x: 1 if x.lower() in stopword_fr else 0)
word_df['is_first_char_upper'] = word_df['word'].apply(lambda x: 1 if x[0].isupper() else 0)
word_df['is_upper'] = word_df['word'].apply(lambda x: 1 if x.isupper() else 0)
word_df['len_word'] = word_df['word'].apply(lambda x: len(x))  # len

# Checking if our random firstname is in french firstname to benchmark
word_df['firstname_is_french'] = 0
word_df.loc[(word_df['is_target'] ==1) & (word_df['word'].isin(firstname_list)), 'firstname_is_french'] = 1

# Add some Random on is_firstname (0/1)
word_df['admin_firstname'] = 0 # Admin stuff 
word_df.loc[(word_df['is_target'] ==1) & (word_df.is_firstname == 1), 'admin_firstname'] = 1 # Admin stuff 
# Random Firstname 2/3 chance to be Firstname
word_df.loc[(word_df['is_target'] ==1) & (word_df.is_firstname == 1), 'is_firstname'] = 1 if random.randint(0, 1) >= 1 else 0 

# Shift is_firstname
word_df['is_firstname_1b'] = word_df.groupby(['doc_name'])['is_firstname'].apply(lambda x: x.shift(-1))
word_df['is_firstname_2b'] = word_df.groupby(['doc_name'])['is_firstname'].apply(lambda x: x.shift(-2))

word_df['is_firstname_1a'] = word_df.groupby(['doc_name'])['is_firstname'].apply(lambda x: x.shift(1))
word_df['is_firstname_2a'] = word_df.groupby(['doc_name'])['is_firstname'].apply(lambda x: x.shift(2))

#Check is there is a "Mr", "M" before...
word_df['is_mister_word'] = 0
word_df.loc[word_df['word'].isin(mister_list), 'is_mister_word'] = 1

## Label encoding word stem
lbl = LabelEncoder()

word_df['word_stem'] = word_df['word'].apply(lambda x: stemmer.stem(x))
word_df['word_encoded'] = lbl.fit_transform(list(word_df['word_stem'].values))
#drop stem word
word_df = word_df.drop('word_stem', axis=1)

# Shift words encoded
## One word before
word_df['word_encoded_shift_1b'] = word_df.groupby(['doc_name'])['word_encoded'].apply(lambda x: x.shift(-1))

## Two words before
word_df['word_encoded_shift_2b'] = word_df.groupby(['doc_name'])['word_encoded'].apply(lambda x: x.shift(-2))

## One word after
word_df['word_encoded_shift_1a'] = word_df.groupby(['doc_name'])['word_encoded'].apply(lambda x: x.shift(1))

## Two words after
word_df['word_encoded_shift_2a'] = word_df.groupby(['doc_name'])['word_encoded'].apply(lambda x: x.shift(2))

# Shift is_mister_word Features
## Beforef 
word_df['is_mister_word_1b'] = word_df.groupby(['doc_name'])['is_mister_word'].apply(lambda x: x.shift(1))
word_df['is_mister_word_2b'] = word_df.groupby(['doc_name'])['is_mister_word'].apply(lambda x: x.shift(2))
## After
word_df['is_mister_word_1a'] = word_df.groupby(['doc_name'])['is_mister_word'].apply(lambda x: x.shift(-1))
word_df['is_mister_word_2a'] = word_df.groupby(['doc_name'])['is_mister_word'].apply(lambda x: x.shift(-2))


# Check type of word
def type_extraction(x):
    """
    Return type of word
    """
    try:
        blob = TextBlob(x, pos_tagger=PatternTagger(), analyzer=PatternAnalyzer())
        return blob.tags[0][1]
    except:
         return 'PONCTUATION'


word_df['word_type'] = word_df['word'].apply(lambda x: type_extraction(x))

# Encode word type
lbl_word_type = LabelEncoder()
word_df['word_type'] = lbl_word_type.fit_transform(list(word_df['word_type'].values))
print " EXPORT WORD TYPE"
word_type_info = pd.DataFrame()
word_type_info['word_type'] = lbl_word_type.classes_
word_type_info['encoded'] = word_type_info.word_type.apply(lambda x: lbl_word_type.transform(unicode(x)))
word_type_info.to_csv('data/word_type_info.csv', encoding=('utf-8'), index=False)

# Shift Word type
word_df['word_type_1b'] = word_df.groupby(['doc_name'])['word_type'].apply(lambda x: x.shift(1))
word_df['word_type_2b'] = word_df.groupby(['doc_name'])['word_type'].apply(lambda x: x.shift(2))
## After
word_df['word_type_1a'] = word_df.groupby(['doc_name'])['word_type'].apply(lambda x: x.shift(-1))
word_df['word_type_2a'] = word_df.groupby(['doc_name'])['word_type'].apply(lambda x: x.shift(-2))

# Fillna Nan word shift
word_df = word_df.fillna(-1)


print " EXPORT BAG OF WORD"
## Export bag of word from LabelEncoding
encoding_label = pd.DataFrame()
encoding_label['words'] = lbl.classes_
encoding_label['encoded'] = encoding_label.words.apply(lambda x: lbl.transform(unicode(x)))
encoding_label.to_csv('data/encoding_label.csv', encoding=('utf-8'), index=False)


# Admin features for benchmark
####################################
# admin_name : random name given for positive tag
# add_row to know if we add the line
# admin_firstname (if it is a firstanme in our base)


print " EXPORT..."
word_df.to_csv('data/data.csv', encoding='utf-8', index=False)


#
# OLD
#docText = '\n\n'.join([
#    paragraph.text.encode('utf-8') for paragraph in paragraphs
#])


# tagger = treetaggerwrapper.TreeTagger(TAGLANG='fr', TAGDIR='/Users/babou/Desktop/NLP',
#  TAGINENC='utf-8',TAGOUTENC='utf-8')
